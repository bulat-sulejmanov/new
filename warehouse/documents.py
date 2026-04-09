from __future__ import annotations

import os
import re
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from django.conf import settings
from django.utils import timezone


@dataclass(frozen=True)
class PurchaseOrderItemDocumentRow:
    index: int
    sku: str
    name: str
    unit: str
    quantity_display: str
    price_display: str
    sum_display: str


@dataclass(frozen=True)
class PurchaseOrderDocumentPayload:
    number: str
    supplier_name: str
    supplier_inn: str
    supplier_kpp: str
    supplier_address: str
    supplier_email: str
    supplier_phone: str
    supplier_contact_person: str
    status_display: str
    expected_date_display: str
    created_at_display: str
    notes: str
    items: tuple[PurchaseOrderItemDocumentRow, ...]
    items_count: int
    total_quantity_display: str
    total_sum_display: str
    company_name: str
    company_full_name: str
    company_inn: str
    company_kpp: str
    company_ogrn: str
    generated_at_display: str
    generated_by_display: str


_DECIMAL_ZERO = Decimal("0")


def _sanitize_export_number(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value or "")
    cleaned = cleaned.strip("._-")
    return cleaned or "purchase-order"


def build_purchase_order_filename(po, extension: str) -> str:
    safe_number = _sanitize_export_number(getattr(po, "number", ""))
    return f"purchase-order-{safe_number}.{extension}"


def _to_decimal(value) -> Decimal | None:
    if value is None or value == "":
        return None
    if isinstance(value, Decimal):
        return value
    try:
        return Decimal(str(value))
    except (InvalidOperation, TypeError, ValueError):
        return None


def _format_decimal(value, *, places: int = 2, strip_trailing_zeros: bool = False) -> str:
    dec = _to_decimal(value)
    if dec is None:
        return "—"

    quant = Decimal("1") if places <= 0 else Decimal("1." + ("0" * places))
    dec = dec.quantize(quant)
    rendered = f"{dec:,.{places}f}"
    rendered = rendered.replace(",", " ").replace(".", ",")
    if strip_trailing_zeros and "," in rendered:
        rendered = rendered.rstrip("0").rstrip(",")
    return rendered


def _format_money(value) -> str:
    return _format_decimal(value, places=2, strip_trailing_zeros=False)


def _format_quantity(value) -> str:
    return _format_decimal(value, places=2, strip_trailing_zeros=True)


def _format_date(value, fmt: str) -> str:
    if not value:
        return "—"
    if hasattr(value, "tzinfo") and value.tzinfo:
        return timezone.localtime(value).strftime(fmt)
    return value.strftime(fmt)


def _iter_po_rows(po) -> Iterable[PurchaseOrderItemDocumentRow]:
    for index, item in enumerate(po.items.select_related("product").all(), start=1):
        product = item.product
        line_sum = item.get_sum()
        yield PurchaseOrderItemDocumentRow(
            index=index,
            sku=product.sku or "—",
            name=product.name or "—",
            unit=product.get_unit_display() or "—",
            quantity_display=_format_quantity(item.quantity),
            price_display=_format_money(item.price) if item.price is not None else "—",
            sum_display=_format_money(line_sum) if line_sum is not None else "—",
        )


def build_purchase_order_payload(po, generated_by=None) -> PurchaseOrderDocumentPayload:
    rows = tuple(_iter_po_rows(po))
    po_items = list(po.items.all())
    total_quantity = sum((_to_decimal(item.quantity) or _DECIMAL_ZERO) for item in po_items)
    total_sum = sum((_to_decimal(item.get_sum()) or _DECIMAL_ZERO) for item in po_items)

    supplier = po.supplier
    generator_name = "Система"
    if generated_by is not None:
        generator_name = generated_by.get_full_name() or generated_by.get_username()

    return PurchaseOrderDocumentPayload(
        number=po.number,
        supplier_name=supplier.name,
        supplier_inn=supplier.inn or "—",
        supplier_kpp=supplier.kpp or "—",
        supplier_address=supplier.address or "—",
        supplier_email=supplier.email or "—",
        supplier_phone=supplier.phone or "—",
        supplier_contact_person=supplier.contact_person or "—",
        status_display=po.get_status_display(),
        expected_date_display=_format_date(po.expected_date, "%d.%m.%Y"),
        created_at_display=_format_date(po.created_at, "%d.%m.%Y %H:%M"),
        notes=po.notes or "",
        items=rows,
        items_count=len(rows),
        total_quantity_display=_format_quantity(total_quantity),
        total_sum_display=_format_money(total_sum),
        company_name=getattr(settings, "COMPANY_NAME", ""),
        company_full_name=getattr(settings, "COMPANY_FULL_NAME", ""),
        company_inn=getattr(settings, "COMPANY_INN", ""),
        company_kpp=getattr(settings, "COMPANY_KPP", ""),
        company_ogrn=getattr(settings, "COMPANY_OGRN", ""),
        generated_at_display=timezone.localtime(timezone.now()).strftime("%d.%m.%Y %H:%M"),
        generated_by_display=generator_name,
    )


def build_purchase_order_docx(po, generated_by=None) -> bytes:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError(
            "Для экспорта в Word установите зависимость python-docx."
        ) from exc

    payload = build_purchase_order_payload(po, generated_by=generated_by)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Заказ поставщику № {payload.number}")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(payload.company_full_name or payload.company_name)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10.5)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Создан: {payload.created_at_display}    |    Статус: {payload.status_display}    |    Ожидаемая дата: {payload.expected_date_display}"
    )

    def add_section_heading(text: str):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        return paragraph

    def add_info_table(rows: list[tuple[str, str]]):
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value or "—"
            for idx, cell in enumerate(cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if paragraph.runs:
                    paragraph.runs[0].font.name = "Arial"
                    paragraph.runs[0].font.size = Pt(10)
                    paragraph.runs[0].bold = idx == 0
        return table

    add_section_heading("Реквизиты заказчика")
    add_info_table(
        [
            ("Организация", payload.company_full_name or payload.company_name or "—"),
            ("ИНН / КПП", f"{payload.company_inn or '—'} / {payload.company_kpp or '—'}"),
            ("ОГРН", payload.company_ogrn or "—"),
        ]
    )

    add_section_heading("Поставщик")
    add_info_table(
        [
            ("Название", payload.supplier_name),
            ("ИНН / КПП", f"{payload.supplier_inn} / {payload.supplier_kpp}"),
            ("Адрес", payload.supplier_address),
            ("Email", payload.supplier_email),
            ("Телефон", payload.supplier_phone),
            ("Контактное лицо", payload.supplier_contact_person),
        ]
    )

    add_section_heading("Позиции заказа")
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["№", "Артикул", "Наименование", "Кол-во", "Ед.", "Цена", "Сумма"]
    widths_cm = [1.0, 3.4, 5.8, 2.2, 1.6, 2.5, 2.7]
    header_cells = table.rows[0].cells
    for cell, text, width_cm in zip(header_cells, headers, widths_cm):
        cell.width = Cm(width_cm)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    if payload.items:
        for row in payload.items:
            cells = table.add_row().cells
            values = [
                str(row.index),
                row.sku,
                row.name,
                row.quantity_display,
                row.unit,
                row.price_display,
                row.sum_display,
            ]
            for idx, (cell, value, width_cm) in enumerate(zip(cells, values, widths_cm)):
                cell.width = Cm(width_cm)
                paragraph = cell.paragraphs[0]
                if idx in {0, 3, 4, 5, 6}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in {0, 4} else WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(value)
                run.font.name = "Arial"
                run.font.size = Pt(10)
    else:
        cells = table.add_row().cells
        cells[0].merge(cells[-1])
        cells[0].text = "Позиции отсутствуют"
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    totals = document.add_paragraph()
    totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_run = totals.add_run(
        f"Итого позиций: {payload.items_count}    |    Общее количество: {payload.total_quantity_display}    |    Общая сумма: {payload.total_sum_display} ₽"
    )
    totals_run.bold = True
    totals_run.font.name = "Arial"
    totals_run.font.size = Pt(10.5)

    add_section_heading("Примечание")
    note_paragraph = document.add_paragraph(payload.notes or "Без примечаний")
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    signatures = document.add_table(rows=2, cols=2)
    signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    signatures.autofit = False
    for idx, width in enumerate([7.8, 7.8]):
        signatures.columns[idx].width = Cm(width)
    signatures.cell(0, 0).text = "Ответственный за заказ"
    signatures.cell(0, 1).text = f"Сформировал: {payload.generated_by_display}"
    signatures.cell(1, 0).text = "______________________________"
    signatures.cell(1, 1).text = "______________________________"
    for row in signatures.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    generated_run = generated.add_run(f"Документ сформирован {payload.generated_at_display}")
    generated_run.italic = True
    generated_run.font.name = "Arial"
    generated_run.font.size = Pt(9)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@lru_cache(maxsize=1)
def _resolve_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    from reportlab.pdfbase.ttfonts import TTFont

    windir = Path(os.environ.get("WINDIR", "C:/Windows"))
    font_candidates = [
        {
            "family": "ArialUnicodeFallback",
            "regular": [
                windir / "Fonts" / "arial.ttf",
                Path("/Library/Fonts/Arial.ttf"),
                Path("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"),
            ],
            "bold": [
                windir / "Fonts" / "arialbd.ttf",
                Path("/Library/Fonts/Arial Bold.ttf"),
                Path("/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf"),
            ],
        },
        {
            "family": "DejaVuSansFallback",
            "regular": [
                Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
                Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
                Path.home() / ".fonts" / "DejaVuSans.ttf",
            ],
            "bold": [
                Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
                Path("/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
                Path.home() / ".fonts" / "DejaVuSans-Bold.ttf",
            ],
        },
        {
            "family": "LiberationSansFallback",
            "regular": [
                Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
                Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
            ],
            "bold": [
                Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
                Path("/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
            ],
        },
    ]

    for candidate in font_candidates:
        regular_path = next((path for path in candidate["regular"] if path.exists()), None)
        if not regular_path:
            continue
        bold_path = next((path for path in candidate["bold"] if path.exists()), regular_path)
        normal_name = f"{candidate['family']}-Regular"
        bold_name = f"{candidate['family']}-Bold"
        if normal_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(normal_name, str(regular_path)))
        if bold_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
        registerFontFamily(candidate["family"], normal=normal_name, bold=bold_name)
        return normal_name, bold_name

    return "Helvetica", "Helvetica-Bold"


def _paragraph_text(value: str) -> str:
    safe_text = escape(value or "—")
    return safe_text.replace("\n", "<br/>")


def build_purchase_order_pdf(po, generated_by=None) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError(
            "Для экспорта в PDF установите зависимости reportlab/xhtml2pdf."
        ) from exc

    regular_font, bold_font = _resolve_pdf_fonts()
    if regular_font == "Helvetica":
        raise RuntimeError(
            "Не найден системный шрифт с поддержкой кириллицы для PDF. Нужен Arial, DejaVu Sans или Liberation Sans."
        )

    payload = build_purchase_order_payload(po, generated_by=generated_by)
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=f"Заказ поставщику № {payload.number}",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "POTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=16,
        leading=20,
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=colors.HexColor("#111827"),
    )
    subtitle_style = ParagraphStyle(
        "POSubtitle",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=9.5,
        leading=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4b5563"),
        spaceAfter=10,
    )
    section_style = ParagraphStyle(
        "POSection",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=11.5,
        leading=14,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "POBody",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=9.3,
        leading=11.5,
        textColor=colors.HexColor("#111827"),
    )
    body_right = ParagraphStyle("POBodyRight", parent=body_style, alignment=TA_RIGHT)
    body_center = ParagraphStyle("POBodyCenter", parent=body_style, alignment=TA_CENTER)
    table_header = ParagraphStyle(
        "POTableHeader",
        parent=body_style,
        fontName=bold_font,
        alignment=TA_CENTER,
        textColor=colors.white,
        fontSize=9,
        leading=10.5,
    )
    small_muted = ParagraphStyle(
        "POSmallMuted",
        parent=body_style,
        fontSize=8.2,
        leading=10,
        textColor=colors.HexColor("#6b7280"),
        alignment=TA_RIGHT,
    )

    def p(text: str, style=body_style):
        return Paragraph(_paragraph_text(text), style)

    story = [
        Paragraph(f"Заказ поставщику № {escape(payload.number)}", title_style),
        Paragraph(
            _paragraph_text(payload.company_full_name or payload.company_name or "Организация не указана"),
            subtitle_style,
        ),
        Table(
            [[p(f"Создан: {payload.created_at_display}"), p(f"Статус: {payload.status_display}"), p(f"Ожидаемая дата: {payload.expected_date_display}")]],
            colWidths=[58 * mm, 52 * mm, 48 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f3f4f6")),
                    ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#d1d5db")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d1d5db")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            ),
        ),
        Spacer(1, 6),
        Paragraph("Реквизиты заказчика", section_style),
        Table(
            [
                [p("Организация"), p(payload.company_full_name or payload.company_name or "—")],
                [p("ИНН / КПП"), p(f"{payload.company_inn or '—'} / {payload.company_kpp or '—'}")],
                [p("ОГРН"), p(payload.company_ogrn or "—")],
            ],
            colWidths=[40 * mm, 134 * mm],
            style=TableStyle(
                [
                    ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cbd5e1")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#e5e7eb")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            ),
        ),
        Spacer(1, 8),
        Paragraph("Поставщик", section_style),
        Table(
            [
                [p("Название"), p(payload.supplier_name)],
                [p("ИНН / КПП"), p(f"{payload.supplier_inn} / {payload.supplier_kpp}")],
                [p("Адрес"), p(payload.supplier_address)],
                [p("Email"), p(payload.supplier_email)],
                [p("Телефон"), p(payload.supplier_phone)],
                [p("Контактное лицо"), p(payload.supplier_contact_person)],
            ],
            colWidths=[40 * mm, 134 * mm],
            style=TableStyle(
                [
                    ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cbd5e1")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#e5e7eb")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            ),
        ),
        Spacer(1, 8),
        Paragraph("Позиции заказа", section_style),
    ]

    table_data = [
        [
            p("№", table_header),
            p("Артикул", table_header),
            p("Наименование", table_header),
            p("Кол-во", table_header),
            p("Ед.", table_header),
            p("Цена", table_header),
            p("Сумма", table_header),
        ]
    ]

    if payload.items:
        for row in payload.items:
            table_data.append(
                [
                    p(str(row.index), body_center),
                    p(row.sku),
                    p(row.name),
                    p(row.quantity_display, body_right),
                    p(row.unit, body_center),
                    p(row.price_display, body_right),
                    p(row.sum_display, body_right),
                ]
            )
    else:
        table_data.append([p("Позиции отсутствуют", body_center)] + [p("", body_center) for _ in range(6)])

    items_table = LongTable(
        table_data,
        colWidths=[11 * mm, 30 * mm, 55 * mm, 18 * mm, 14 * mm, 22 * mm, 24 * mm],
        repeatRows=1,
    )
    items_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cbd5e1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#e5e7eb")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
            ]
        )
    )
    story.append(items_table)
    story.extend(
        [
            Spacer(1, 6),
            Table(
                [[p(f"Итого позиций: {payload.items_count}    Общее количество: {payload.total_quantity_display}    Общая сумма: {payload.total_sum_display} ₽", body_right)]],
                colWidths=[174 * mm],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eff6ff")),
                        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#bfdbfe")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                ),
            ),
            Spacer(1, 8),
            Paragraph("Примечание", section_style),
            Table(
                [[p(payload.notes or "Без примечаний")]],
                colWidths=[174 * mm],
                style=TableStyle(
                    [
                        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#d1d5db")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                ),
            ),
            Spacer(1, 12),
            Table(
                [
                    [p("Ответственный за заказ"), p(f"Сформировал: {payload.generated_by_display}")],
                    [p("______________________________"), p("______________________________")],
                ],
                colWidths=[87 * mm, 87 * mm],
                style=TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ]
                ),
            ),
            Spacer(1, 10),
            Paragraph(f"Документ сформирован {payload.generated_at_display}", small_muted),
        ]
    )

    document.build(story)
    return buffer.getvalue()
